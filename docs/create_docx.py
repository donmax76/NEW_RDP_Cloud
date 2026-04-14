from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Page margins
for s in doc.sections:
    s.top_margin = Cm(2); s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prometey")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Screenshot List for PDF Guide")
r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("20 screenshots required from the client interface")
r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_paragraph()

# General Rules
h = doc.add_heading("General Rules", level=1)
rules = [
    ("Format:", "PNG (not JPEG)"),
    ("Browser:", "Chrome / Edge, language EN, window width 1400+ px"),
    ("Private data:", "Mask passwords and IP addresses with a gray rectangle"),
    ("OS scale:", "Set to 100% during capture for crisp text"),
    ("Save to:", "docs/img/ with exact filenames from the tables below"),
]
for label, text in rules:
    p = doc.add_paragraph()
    r = p.add_run(label + " "); r.bold = True; r.font.size = Pt(11)
    r = p.add_run(text); r.font.size = Pt(11)

doc.add_paragraph()

def add_table(title, rows):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.rows[0]
    for i, txt in enumerate(["#", "Filename", "What to capture", "Size"]):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Dark header background
        shading = cell._element.get_or_add_tcPr()
        sh = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2E4057', qn('w:val'): 'clear', qn('w:color'): 'auto'})
        shading.append(sh)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, txt in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            r.font.size = Pt(10)
            if ci <= 1: r.bold = True
            # Alternating rows
            if ri % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                sh = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F0F4F8', qn('w:val'): 'clear', qn('w:color'): 'auto'})
                shading.append(sh)

    # Column widths
    widths = [Cm(1.2), Cm(4.5), Cm(8.5), Cm(2.5)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w

    doc.add_paragraph()

# Table 1
add_table("Connection (3 screenshots)", [
    ["1", "01_login.png", "Login page: Token + Password fields empty, Connect button visible", "900x600"],
    ["2", "02_ssl_warning.png", 'Chrome "Your connection is not private" page (incognito)', "1200x700"],
    ["3", "03_connected.png", "Topbar after connecting: green dot, Host: ONLINE, versions on the right", "1600x100"],
])

# Table 2
add_table("Panels (12 screenshots)", [
    ["4", "04_dashboard.png", "Full Dashboard: RAM/CPU/GPU/Uptime cards + Activity Log + Quick Actions", "1600x1000"],
    ["5", "05_speed_test.png", "Speed Test block after Run All Tests - all 4 cards with colored values", "1600x500"],
    ["6", "06_screen.png", "Screen panel with active remote desktop video stream", "1600x1000"],
    ["7", "07_files.png", "Files: disk/folder tree on left + file list with columns on right", "1600x1000"],
    ["8", "08_processes.png", "Processes: table with PID / Name / CPU% / Memory / Threads columns", "1600x800"],
    ["9", "09_terminal.png", "Terminal panel after running ipconfig command", "1600x700"],
    ["10", "10_audio.png", "Audio Recording: recording list on left + waveform player on right", "1600x700"],
    ["11", "11_screenshots.png", "Screenshots: grid view with 4-6 thumbnail previews", "1600x800"],
    ["12", "12_eventlog.png", "Event Log with entries + Auto-clean toolbar (Once / Loop / Off)", "1600x800"],
    ["13", "13_services.png", "Services: list of Windows services with Name / Status / Start Type", "1600x800"],
    ["14", "14_programs.png", "Installed Programs: table with Name / Version / Publisher / Size", "1600x800"],
    ["15", "15_registry.png", "Registry Editor: key tree on left + values on right", "1600x800"],
])

# Table 3
add_table("Settings (5 screenshots)", [
    ["16", "16_settings_update.png", "Block Remote Host Update: DLL URL + Update Host + Restart buttons", "1600x400"],
    ["17", "17_settings_deploy.png", "Block VPS Deploy: 3 file inputs + green Upload All 3 button", "1600x400"],
    ["18", "18_settings_config.png", "Host Config Editor: loaded JSON + New / Open / Save buttons", "1600x500"],
    ["19", "19_settings_threat.png", "Threat Monitor block: checkboxes + orange border + Check now button", "1600x300"],
    ["20", "20_lang_switch.png", "Language selector in topbar OPEN - dropdown showing EN / RU / AZ", "400x200"],
])

# Note
p = doc.add_paragraph()
r = p.add_run("! Note: ")
r.bold = True; r.font.color.rgb = RGBColor(0xE6, 0x51, 0x00); r.font.size = Pt(11)
r = p.add_run("VPS terminal and Windows host installation screenshots will be generated synthetically. Only the 20 client UI screenshots above need to be captured manually.")
r.font.color.rgb = RGBColor(0xE6, 0x51, 0x00); r.font.size = Pt(11)

out = "D:/Android_Projects/NEW_RDP_Cloud/docs/screenshots_list.docx"
doc.save(out)
print(f"Created: {out}")
