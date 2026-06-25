"""Generate Word documents from PROMETEY presentation content (RU / EN / AZ)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ACCENT = RGBColor(0x25, 0x63, 0xEB)   # blue
GREEN  = RGBColor(0x16, 0xA3, 0x4A)
GRAY   = RGBColor(0x64, 0x74, 0x8B)
BLACK  = RGBColor(0x0F, 0x17, 0x2A)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_section_heading(doc, num, title, lang='ru'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run_num = p.add_run(num + '  ')
    run_num.font.size = Pt(9)
    run_num.font.color.rgb = GRAY
    run_num.font.bold = False
    run_title = p.add_run(title)
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = BLACK
    p.paragraph_format.keep_with_next = True

def add_desc(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = GRAY

def add_card_grid(doc, cards):
    """cards = list of (icon, title, desc)"""
    for icon, title, desc in cards:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.keep_together = True
        r = p.add_run(f'{icon}  {title}')
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = ACCENT
        p2 = doc.add_paragraph(desc)
        p2.paragraph_format.left_indent = Cm(1.2)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        p2.runs[0].font.size = Pt(10)
        p2.runs[0].font.color.rgb = GRAY

def add_enc_card(doc, title, body_paras, tech):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.keep_together = True
    r = p.add_run(title)
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLACK
    for text in body_paras:
        p2 = doc.add_paragraph(text)
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        p2.runs[0].font.size = Pt(10)
    p3 = doc.add_paragraph(tech)
    p3.paragraph_format.left_indent = Cm(0.5)
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(8)
    p3.runs[0].font.size = Pt(9)
    p3.runs[0].font.color.rgb = ACCENT
    p3.runs[0].font.italic = True

def add_layer(doc, num, title, plain, tech=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.keep_together = True
    r = p.add_run(f'{num}.  {title}')
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLACK
    p2 = doc.add_paragraph(plain)
    p2.paragraph_format.left_indent = Cm(0.8)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    p2.runs[0].font.size = Pt(10)
    if tech:
        p3 = doc.add_paragraph(tech)
        p3.paragraph_format.left_indent = Cm(0.8)
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(8)
        p3.runs[0].font.size = Pt(9)
        p3.runs[0].font.color.rgb = ACCENT
        p3.runs[0].font.italic = True

def add_infobox(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = GRAY

def add_av_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], '1E3A8A')
        run = hdr[i].paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(9)
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = val
            cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# RUSSIAN
# ─────────────────────────────────────────────────────────────────────────────
def build_ru():
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run('⚡ DATA')
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    p2 = doc.add_paragraph('Система скрытого удалённого управления')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(14)
    p2.runs[0].font.color.rgb = GRAY

    doc.add_paragraph('v1.0.250  ·  2026').paragraphs[-1] if False else None
    pv = doc.add_paragraph('v1.0.250  ·  2026')
    pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pv.runs[0].font.size = Pt(10)
    pv.runs[0].font.color.rgb = GRAY

    doc.add_page_break()

    # ── Section 1: Architecture ──────────────────────────────────────────────
    add_section_heading(doc, 'Раздел 1', 'Архитектура системы')
    add_desc(doc, 'Четыре уровня: Объект → VPS-1 + Cloudflare → VPS-2 (relay) → Оператор (браузер). '
             'Реальный IP никогда не раскрывается. Цепочка двух VPS + Cloudflare proxy.')

    arch_items = [
        ('🖥️', 'Объект (управляемый компьютер)',
         'Невидимая служба · агент внутри svchost.exe · '
         'автозапуск с Windows · конфиг зашифрован AES-256'),
        ('☁️', 'VPS-1 + Cloudflare (домен + CDN)',
         'Точка подключения объекта · wss://domain.com · реальный IP скрыт за Cloudflare · '
         'DDoS-защита · бесплатный SSL сертификат CF'),
        ('☁️', 'VPS-2 — Relay (скрытый IP, без CF)',
         'Мост между объектом и оператором · IP знает только VPS-1 · журналы только в RAM · без Cloudflare'),
        ('👨‍💻', 'Оператор',
         'Управляет в браузере с любого устройства · вход: логин + пароль · роли: Администратор / Оператор · '
         'IP объекту не виден'),
    ]
    add_card_grid(doc, arch_items)

    add_infobox(doc,
        '🔍 Поток данных: Оператор → TLS 1.3 → VPS-2 → VPS-1 → Cloudflare → Объект. '
        'Весь канал выглядит как обычный HTTPS/CDN-трафик.')

    add_infobox(doc,
        '🌐 Преимущества 2 VPS + Cloudflare: доменное имя вместо IP, бесплатный SSL, '
        'реальный IP VPS-1 за Cloudflare, изоляция VPS-2, DDoS-защита, разные юрисдикции.')

    doc.add_page_break()

    # ── Section 2: Capabilities ──────────────────────────────────────────────
    add_section_heading(doc, 'Раздел 2', 'Возможности системы')
    add_desc(doc, 'Все функции работают в браузере. На компьютер оператора не устанавливается никаких программ.')

    caps = [
        ('🖥️', 'Трансляция экрана в реальном времени',
         'H.264 (QVBR) и JPEG. Регулируются FPS, качество, масштаб, битрейт. WSS/TCP порт 443.'),
        ('📁', 'Файловый менеджер',
         'Доступ ко всем папкам включая системные. Загрузка, скачивание, удаление, переименование.'),
        ('⚙️', 'Процессы и службы',
         'Все запущенные программы и службы. Завершение, запуск, изменение типа службы.'),
        ('💻', 'Терминал (cmd / PowerShell)',
         'Полный терминал с правами SYSTEM. PowerShell и cmd — результат немедленно.'),
        ('📋', 'Реестр Windows',
         'Создание, редактирование, удаление ключей и значений. Все типы данных.'),
        ('📷', 'Автоматические скриншоты',
         'Снимки через заданные интервалы. Для определённых программ или постоянно. Хранение на VPS.'),
        ('🎙️', 'Аудио: запись и прослушивание',
         'Микрофон и системный звук. WASAPI loopback — индикатор микрофона не отображается. Кодек Opus.'),
        ('📈', 'История активности',
         'Включение/выключение/блокировка. Общее время работы, количество сессий.'),
        ('👤', 'Управление операторами',
         'Несколько операторов с разными правами. Каждый входит со своим логином.'),
        ('🔄', 'Удалённое обновление',
         'Загрузите новую версию агента из браузера. Физический доступ не нужен.'),
        ('🛡️', 'Управление защитой',
         'Windows Defender, очистка системных журналов, удаление следов активности.'),
        ('🔥', 'Самоуничтожение',
         'Полное удаление одной кнопкой: служба, файлы, реестр. Ничего не остаётся.'),
        ('📍', 'WPS Геолокация',
         'Координаты объекта через Windows Location API (GPS / Wi-Fi WPS). Точность до нескольких метров. Список ближайших Wi-Fi сетей. Автообновление.'),
    ]
    add_card_grid(doc, caps)

    add_infobox(doc,
        '⚙️ Системные требования:  Объект: Windows 11 x64  |  Оператор: любая ОС, достаточно браузера')

    doc.add_page_break()

    # ── Section 3: Encryption ────────────────────────────────────────────────
    add_section_heading(doc, 'Раздел 3', 'Шифрование и защита канала')
    add_desc(doc, 'Четыре независимых уровня криптографической защиты. '
             'Каждый байт трафика защищён — от передачи до хранения.')

    add_enc_card(doc, '🌐  TLS 1.3 — Транспортное шифрование',
        ['Весь трафик Object ↔ VPS ↔ Operator шифруется TLS 1.3 — современнейший стандарт. '
         'Ephemeral Diffie-Hellman: сессионные ключи только в RAM, никогда не передаются по сети (Forward Secrecy). '
         'Снифер видит только случайные байты.'],
        'TLS_AES_256_GCM_SHA384 · X25519 ECDHE · Port 443 · nginx reverse proxy')

    add_enc_card(doc, '🔒  WSS — WebSocket Secure',
        ['Все каналы — WSS (WebSocket over HTTPS). Постоянный двусторонний канал поверх HTTPS. '
         'С внешней стороны трафик неотличим от обычного HTTPS-сайта. DPI не определит.'],
        'wss://server:443/host · wss://server:443/ws · Upgrade: websocket')

    add_enc_card(doc, '🔑  AES-256-GCM — Шифрование конфигурации',
        ['Конфиг-файл на диске зашифрован AES-256-GCM. '
         'Защищает: IP ВПС, токен комнаты, пароль. Без агента файл нечитаем. '
         'GCM = аутентифицированное шифрование (AEAD) — изменение байта → отклонение блока. '
         '2²⁵⁶ комбинаций — подбор невозможен даже за миллиарды лет.'],
        'AES-256-GCM · 96-bit IV · 128-bit auth tag')

    add_enc_card(doc, '💾  AES-256-CBC — Шифрование конфига (резервный слой)',
        ['Конфиг-файл: ключ извлекается из пароля PBKDF2-HMAC-SHA256 (100 000 итераций), '
         'затем AES-256-CBC. Физический анализ диска не поможет — ключ только в памяти агента.'],
        'AES-256-CBC · PBKDF2-HMAC-SHA256 (100K iter) · random salt · random IV')

    add_infobox(doc,
        '💡 Forward Secrecy: для каждой TLS-сессии создаются временные DH-ключи. '
        'Даже при краже серверного сертификата расшифровать прошлые сессии невозможно.')

    doc.add_page_break()

    # ── Section 3b: Network Anonymity ────────────────────────────────────────
    add_section_heading(doc, 'Раздел 3 (продолжение)', 'Анонимность сети')
    add_desc(doc, 'Оператор никогда не раскрывается. VPS в другой стране. '
             'Следователи видят только обычный HTTPS-трафик.')

    doc.add_paragraph('❌  Что видит снифер / следователь:').runs[0].font.bold = True
    sniffer_items = [
        'IP сервера — только общие IP Cloudflare (104.x.x.x). Реальный IP VPS-1 никогда не виден.',
        'Содержимое трафика — TLS 1.3. Даже Cloudflare не видит внутренний трафик.',
        'Тип активности — обычный HTTPS/CDN. Отличить от посещения сайта невозможно.',
        'IP оператора — объект никогда не видит. IP VPS-2 знает только VPS-1.',
    ]
    for item in sniffer_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.5)
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph('✅  Как защищают Cloudflare + два VPS:').runs[0].font.bold = True
    protect_items = [
        'Реальный IP скрыт — VPS-1 за Cloudflare. Три юрисдикции.',
        'Изоляция VPS-2 — IP нигде публично не фигурирует.',
        'Журналов нет — только в RAM. После перезапуска следов нет.',
        'Пароли не раскрываются — хэш PBKDF2.',
    ]
    for item in protect_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.5)
        p.runs[0].font.size = Pt(10)

    add_infobox(doc,
        '🔵 Максимум что можно узнать: объект подключается к domain.com (Cloudflare) порт 443 — '
        'неотличимо от любого сайта. Без одновременных международных запросов к Cloudflare и обоим '
        'VPS-провайдерам — ни цель, ни оператор, ни содержимое не определяются.')

    doc.add_page_break()

    # ── Section 4: Security Layers ───────────────────────────────────────────
    add_section_heading(doc, 'Раздел 4', 'Уровни защиты')
    add_desc(doc, 'Семь независимых уровней защиты со всех сторон — от сетевого до физического анализа диска.')

    add_layer(doc, 1, '🌐 Сетевое шифрование — TLS 1.3',
        'Весь трафик TLS 1.3, порт 443. Forward Secrecy — прошлые сессии нельзя расшифровать.',
        'WebSocket Secure (WSS) · TLS_AES_256_GCM_SHA384 · X25519 · ECDSA')

    add_layer(doc, 2, '🔓 Защита паролем — PBKDF2',
        'Пароли: PBKDF2-HMAC-SHA256 (100 000 итераций) + per-user salt. Подбор займёт годы.',
        'Session token — одноразовый, только в RAM, теряется при перезапуске')

    add_layer(doc, 3, '🧬 Никаких следов на диске и в памяти',
        'Только два файла: агент + зашифрованный конфиг. Никаких логов, кэша, временных файлов.',
        'TLS + AES-256-GCM — трафик зашифрован при передаче; конфиг зашифрован на диске')

    add_layer(doc, 4, '🚫 Невидимость для антивирусов',
        'VirusTotal: Elastic, THOR YARA, Windows Defender — ни один не обнаруживает.',
        '/DELAYLOAD 17 DLL · Microsoft metadata · String obfuscation')

    add_layer(doc, 5, '👻 Полная невидимость для пользователя',
        'Нет иконки, окна, записи в трее, индикатора микрофона. Маскируется под системный компонент.',
        'WASAPI из контекста SYSTEM · DACL + failureflag=1 (не остановить даже админу)')

    add_layer(doc, 6, '💾 Шифрование конфига',
        'Конфиг-файл: AES-256-CBC + PBKDF2. Сервер, пароль, токен — ничего в открытом виде.')

    add_layer(doc, 7, '🧹 Самоочистка',
        'При остановке: модули выгружаются, временные файлы удаляются, системные журналы очищаются.')

    doc.add_page_break()

    # ── Section 5: AV Results ────────────────────────────────────────────────
    add_section_heading(doc, 'Раздел 5', 'Результаты антивирусного теста')
    add_desc(doc, 'Агентский модуль протестирован на VirusTotal.')

    add_av_table(doc,
        ['Движок', 'Статус', 'Как достигнуто'],
        [
            ['Elastic ML',       '✓ Не обнаружено', 'Microsoft metadata + /DELAYLOAD 17 DLL + очистка следов OpenSSL'],
            ['THOR YARA',        '✓ Не обнаружено', 'Строки собираются из фрагментов во время выполнения — YARA не совпадают'],
            ['Windows Defender', '✓ Не обнаружено', 'Цифровая подпись + Microsoft metadata + отложенная загрузка DLL'],
            ['T1057 Process',    '✓ Снижено',        'CreateToolhelp32Snapshot вместо NtQuerySystemInformation — нет следов в IAT'],
        ]
    )

    add_infobox(doc,
        'Примечание: Поведенческие метки в sandbox VirusTotal (T1027, T1071) — информационные аннотации, '
        'не обнаружение. Ни один AV-движок не вынес вердикт "Malicious".')

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# ENGLISH
# ─────────────────────────────────────────────────────────────────────────────
def build_en():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run('⚡ DATA')
    r.font.size = Pt(36); r.font.bold = True; r.font.color.rgb = ACCENT

    p2 = doc.add_paragraph('Covert Remote Control System')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(14); p2.runs[0].font.color.rgb = GRAY

    pv = doc.add_paragraph('v1.0.250  ·  2026')
    pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pv.runs[0].font.size = Pt(10); pv.runs[0].font.color.rgb = GRAY

    doc.add_page_break()

    add_section_heading(doc, 'Section 1', 'System Architecture')
    add_desc(doc, 'Four layers: Object → VPS-1 + Cloudflare → VPS-2 (relay) → Operator (browser). '
             'Real IP never exposed. Dual-VPS + Cloudflare proxy chain.')

    add_card_grid(doc, [
        ('🖥️', 'Object (managed computer)',
         'Invisible service · agent inside svchost.exe · auto-start · AES-256 encrypted config'),
        ('☁️', 'VPS-1 + Cloudflare (domain + CDN)',
         'Object connection point · wss://domain.com · real IP hidden behind Cloudflare · DDoS protection · free SSL'),
        ('☁️', 'VPS-2 — Relay (hidden IP, no CF)',
         'Bridge between object and operator · IP known only to VPS-1 · logs only in RAM'),
        ('👨‍💻', 'Operator',
         'Controls via browser from any device · login + password · roles: Admin / Operator · IP never exposed to object'),
    ])

    add_infobox(doc,
        '🔍 Data flow: Operator → TLS 1.3 → VPS-2 → VPS-1 → Cloudflare → Object. '
        'Entire channel looks like normal HTTPS/CDN traffic.')

    doc.add_page_break()

    add_section_heading(doc, 'Section 2', 'System Capabilities')
    add_desc(doc, 'All features run in the browser. No software installed on the operator\'s computer.')

    add_card_grid(doc, [
        ('🖥️', 'Real-time Screen Streaming',       'H.264 (QVBR) and JPEG. FPS, quality, scale, bitrate controls. WSS/TCP port 443.'),
        ('📁', 'File Manager',                      'Access to all folders including system. Upload, download, delete, rename.'),
        ('⚙️', 'Processes & Services',              'All running programs and services. Terminate, launch, change service type.'),
        ('💻', 'Terminal (cmd / PowerShell)',        'Full terminal with SYSTEM privileges. PowerShell and cmd — instant output.'),
        ('📋', 'Windows Registry',                  'Create, edit, delete keys and values. All data types supported.'),
        ('📷', 'Automatic Screenshots',             'Scheduled captures. For specific apps or continuously. Stored on VPS.'),
        ('🎙️', 'Audio: Record & Listen',           'Mic and system audio. WASAPI loopback — no mic indicator shown. Opus codec.'),
        ('📈', 'Activity History',                  'Power on/off/lock tracking. Total uptime, session count.'),
        ('👤', 'Operator Management',               'Multiple operators with different permissions. Each with own login.'),
        ('🔄', 'Remote Update',                     'Upload new agent version from browser. No physical access needed.'),
        ('🛡️', 'Defense Management',               'Windows Defender, system log cleanup, activity trace deletion.'),
        ('🔥', 'Self-Destruction',                  'Complete removal: service stopped, files deleted, registry cleaned. Nothing remains.'),
        ('📍', 'WPS Geolocation',                   'Object coordinates via Windows Location API (GPS / Wi-Fi WPS). Accuracy to a few meters. Nearby Wi-Fi networks list. Auto-refresh.'),
    ])

    add_infobox(doc, '⚙️ Requirements:  Object: Windows 11 x64  |  Operator: any OS, browser only')

    doc.add_page_break()

    add_section_heading(doc, 'Section 3', 'Encryption & Channel Security')
    add_desc(doc, 'Four independent cryptographic layers. Every byte is protected — in transit and at rest.')

    add_enc_card(doc, '🌐  TLS 1.3 — Transport Encryption',
        ['All traffic Object ↔ VPS ↔ Operator encrypted with TLS 1.3. '
         'Ephemeral DH: session keys only in RAM, never transmitted (Forward Secrecy). '
         'Sniffer sees only random bytes.'],
        'TLS_AES_256_GCM_SHA384 · X25519 ECDHE · Port 443 · nginx reverse proxy')

    add_enc_card(doc, '🔒  WSS — WebSocket Secure',
        ['All channels use WSS (WebSocket over HTTPS). Persistent bidirectional channel over HTTPS. '
         'Traffic indistinguishable from normal HTTPS. DPI cannot identify it.'],
        'wss://server:443/host · wss://server:443/ws · Upgrade: websocket')

    add_enc_card(doc, '🔑  AES-256-GCM — Config Encryption',
        ['Config file on disk encrypted with AES-256-GCM. '
         'Protects: VPS IP, room token, password. Without the agent the file cannot be read. '
         'GCM = AEAD — any byte change → block rejected. 2²⁵⁶ combinations.'],
        'AES-256-GCM · 96-bit IV · 128-bit auth tag')

    add_enc_card(doc, '💾  AES-256-CBC — Config Layer',
        ['Config file: key derived from password via PBKDF2-HMAC-SHA256 (100K iterations), '
         'then AES-256-CBC. Physical disk analysis useless — key only in agent memory.'],
        'AES-256-CBC · PBKDF2-HMAC-SHA256 (100K iter) · random salt · random IV')

    doc.add_page_break()

    add_section_heading(doc, 'Section 3 (cont.)', 'Network Anonymity')
    add_desc(doc, 'Operator is never exposed. VPS in another country. '
             'Investigators see only normal HTTPS traffic.')

    add_infobox(doc,
        '🔵 Maximum discoverable: object connects to domain.com (Cloudflare) port 443 — '
        'indistinguishable from any website. Without simultaneous international requests to Cloudflare '
        'and both VPS providers — no target, operator, or content can be identified.')

    doc.add_page_break()

    add_section_heading(doc, 'Section 4', 'Security Layers')
    add_desc(doc, 'Seven independent security layers covering every attack surface.')

    for num, title, plain, tech in [
        (1, '🌐 Network Encryption — TLS 1.3',
         'All traffic TLS 1.3, port 443. Forward Secrecy.',
         'WSS · TLS_AES_256_GCM_SHA384 · X25519 · ECDSA'),
        (2, '🔓 Password Protection — PBKDF2',
         'PBKDF2-HMAC-SHA256 (100K iterations) + per-user salt. Brute-force takes years.',
         'Session token — single-use, RAM only, expires on restart'),
        (3, '🧬 No Disk/Memory Traces',
         'Only two files: agent + encrypted config. No logs, no cache, no temp files.',
         'TLS + AES-256-GCM — in-transit and at-rest encryption'),
        (4, '🚫 AV Invisibility',
         'VirusTotal: Elastic, THOR YARA, Windows Defender — none detect.',
         '/DELAYLOAD 17 DLL · Microsoft metadata · String obfuscation'),
        (5, '👻 Full User Invisibility',
         'No icon, no window, no tray entry, no mic indicator.',
         'WASAPI from SYSTEM context · DACL + failureflag=1'),
        (6, '💾 Config Encryption',
         'Config file: AES-256-CBC + PBKDF2. Nothing in plaintext.', None),
        (7, '🧹 Self-Cleanup',
         'On stop: modules unloaded, temp files deleted, event logs cleared.', None),
    ]:
        add_layer(doc, num, title, plain, tech)

    doc.add_page_break()

    add_section_heading(doc, 'Section 5', 'Antivirus Test Results')
    add_desc(doc, 'Agent module tested on VirusTotal.')

    add_av_table(doc,
        ['Engine', 'Status', 'How achieved'],
        [
            ['Elastic ML',       '✓ Not detected', 'Microsoft metadata + /DELAYLOAD 17 DLL + OpenSSL string scrubbing'],
            ['THOR YARA',        '✓ Not detected', 'Suspicious strings assembled from fragments at runtime — YARA no match'],
            ['Windows Defender', '✓ Not detected', 'Digital signature + Microsoft metadata + delayed DLL loading'],
            ['T1057 Process',    '✓ Reduced',       'CreateToolhelp32Snapshot instead of NtQuerySystemInformation — no IAT trace'],
        ]
    )

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# AZERBAIJANI
# ─────────────────────────────────────────────────────────────────────────────
def build_az():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run('⚡ DATA')
    r.font.size = Pt(36); r.font.bold = True; r.font.color.rgb = ACCENT

    p2 = doc.add_paragraph('Gizli Uzaqdan İdarəetmə Sistemi')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(14); p2.runs[0].font.color.rgb = GRAY

    pv = doc.add_paragraph('v1.0.250  ·  2026')
    pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pv.runs[0].font.size = Pt(10); pv.runs[0].font.color.rgb = GRAY

    doc.add_page_break()

    add_section_heading(doc, 'Hissə 1', 'Sistem Arxitekturası')
    add_desc(doc, 'Dörd səviyyə: Obyekt → VPS-1 + Cloudflare → VPS-2 (relay) → Operator (brauzer). '
             'Real IP heç vaxt açıqlanmır.')

    add_card_grid(doc, [
        ('🖥️', 'Obyekt (idarə olunan kompüter)',
         'Görünməz servis · agent svchost.exe daxilində · avtoyükləmə · AES-256 şifrəli konfiq'),
        ('☁️', 'VPS-1 + Cloudflare (domen + CDN)',
         'Obyektin qoşulma nöqtəsi · wss://domain.com · real IP Cloudflare arxasında gizlidir · pulsuz SSL'),
        ('☁️', 'VPS-2 — Relay (gizli IP, CF yoxdur)',
         'Obyekt və operator arasında körpü · IP yalnız VPS-1-ə məlumdur · jurnallar yalnız RAM-da'),
        ('👨‍💻', 'Operator',
         'İstənilən cihazdan brauzerdə idarə edir · giriş: login + şifrə · rollar: Admin / Operator'),
    ])

    doc.add_page_break()

    add_section_heading(doc, 'Hissə 2', 'Sistem İmkanları')
    add_desc(doc, 'Bütün funksiyalar brauzerdə işləyir. Operatorun kompüterinə heç bir proqram quraşdırılmır.')

    add_card_grid(doc, [
        ('🖥️', 'Real vaxt ekran yayımı',         'H.264 (QVBR) və JPEG. FPS, keyfiyyət, miqyas, bitreyt idarəsi. WSS/TCP 443 port.'),
        ('📁', 'Fayl meneceri',                   'Sistem qovluqları da daxil olmaqla hamısına giriş. Yükləmə, endirmə, silmə, adlandırma.'),
        ('⚙️', 'Proseslər və xidmətlər',          'Bütün işləyən proqramlar. Sonlandırma, başlatma, xidmət növünü dəyişmə.'),
        ('💻', 'Terminal (cmd / PowerShell)',      'SYSTEM səlahiyyətləri ilə tam terminal. Ani nəticə.'),
        ('📋', 'Windows Reyestri',                'Açar və dəyərlərin yaradılması, redaktəsi, silinməsi.'),
        ('📷', 'Avtomatik ekran görüntüləri',     'Müəyyən intervallarda. VPS-də saxlanılır.'),
        ('🎙️', 'Audio: qeyd və dinləmə',         'Mikrofon və sistem səsi. WASAPI loopback — mikrofon indikatoru görünmür. Opus kodeki.'),
        ('📈', 'Aktivlik tarixi',                 'Açılma/söndürülmə/kilidlənmə. Ümumi iş vaxtı, sessiya sayı.'),
        ('👤', 'Operator idarəetməsi',            'Müxtəlif icazələrlə bir neçə operator. Hər biri öz loginlə daxil olur.'),
        ('🔄', 'Uzaqdan yeniləmə',               'Brauzerdən yeni agent versiyasını yükləyin. Fiziki giriş lazım deyil.'),
        ('🛡️', 'Mühafizə idarəetməsi',          'Windows Defender, sistem jurnallarının təmizlənməsi, aktivlik izlərinin silinməsi.'),
        ('🔥', 'Öz-özünü məhvetmə',             'Bir düyməylə tam silinmə: servis, fayllar, reyestr. Heç nə qalmır.'),
        ('📍', 'WPS Geolokasiya',               'Windows Location API (GPS / Wi-Fi WPS) vasitəsilə koordinatlar. Bir neçə metr dəqiqliyi. Yaxın Wi-Fi siyahısı. Avtoyeniləmə.'),
    ])

    add_infobox(doc, '⚙️ Sistem Tələbləri:  Obyekt: Windows 11 x64  |  Operator: İstənilən OS, brauzer kifayətdir')

    doc.add_page_break()

    add_section_heading(doc, 'Hissə 3', 'Şifrələmə və Kanal Müdafiəsi')
    add_desc(doc, 'Dörd müstəqil kriptoqrafik qoruma səviyyəsi.')

    add_enc_card(doc, '🌐  TLS 1.3 — Nəqliyyat Şifrələməsi',
        ['Bütün trafik TLS 1.3 ilə şifrələnir. Ephemeral DH: sessiya açarları yalnız RAM-da (Forward Secrecy). '
         'Sniffer yalnız təsadüfi baytlar görür.'],
        'TLS_AES_256_GCM_SHA384 · X25519 ECDHE · 443 Port · nginx')

    add_enc_card(doc, '🔑  AES-256-GCM — Konfiqurasiya Şifrələməsi',
        ['Konfiq-faylın diski AES-256-GCM ilə şifrələnir. '
         'Qoruyur: VPS IP, otaq tokeni, şifrə. Agentsiz fayl oxunmaz. AEAD — bir bayt dəyişsə blok rədd edilir.'],
        'AES-256-GCM · 96-bit IV · 128-bit auth tag')

    doc.add_page_break()

    add_section_heading(doc, 'Hissə 4', 'Qoruma Səviyyələri')
    add_desc(doc, 'Yeddi müstəqil qoruma səviyyəsi.')

    for num, title, plain in [
        (1, '🌐 Şəbəkə Şifrələməsi — TLS 1.3',    'Bütün trafik TLS 1.3, 443 port. Forward Secrecy.'),
        (2, '🔓 Şifrə Mühafizəsi — PBKDF2',       'PBKDF2-HMAC-SHA256 (100K iterasiya) + per-user salt.'),
        (3, '🧬 Diskdə iz yoxdur',                 'Yalnız iki fayl: agent + şifrəli konfiq. Log, keş, müvəqqəti fayl yoxdur.'),
        (4, '🚫 Antiviruslara görünməzlik',        'VirusTotal: Elastic, THOR YARA, Windows Defender — heç biri aşkar etmir.'),
        (5, '👻 İstifadəçiyə tam görünməzlik',     'İkon, pəncərə, sistem tepsisi qeydi, mikrofon indikatoru yoxdur.'),
        (6, '💾 Konfiq Şifrələməsi',              'Konfiq-fayl: AES-256-CBC + PBKDF2. Açıq mətn yoxdur.'),
        (7, '🧹 Özünü Təmizləmə',                'Dayananda: modullar boşaldılır, müvəqqəti fayllar silinir, jurnallar təmizlənir.'),
    ]:
        add_layer(doc, num, title, plain)

    doc.add_page_break()

    add_section_heading(doc, 'Hissə 5', 'Antivirus Test Nəticələri')
    add_desc(doc, 'Agent modulu VirusTotal-da test edilib.')

    add_av_table(doc,
        ['Mühərrik', 'Status', 'Necə əldə edilib'],
        [
            ['Elastic ML',       '✓ Aşkar edilmədi', 'Microsoft metadata + /DELAYLOAD 17 DLL'],
            ['THOR YARA',        '✓ Aşkar edilmədi', 'Şübhəli stringlər runtime-da fraqmentlərdən yığılır'],
            ['Windows Defender', '✓ Aşkar edilmədi', 'Rəqəmsal imza + Microsoft metadata'],
            ['T1057 Process',    '✓ Azaldıldı',      'CreateToolhelp32Snapshot — IAT-da iz yoxdur'],
        ]
    )

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    import os
    base = os.path.dirname(os.path.abspath(__file__))

    docs = [
        (build_ru(), 'PROMETEY_RU.docx'),
        (build_en(), 'PROMETEY_EN.docx'),
        (build_az(), 'PROMETEY_AZ.docx'),
    ]
    for doc, name in docs:
        path = os.path.join(base, name)
        doc.save(path)
        print(f'Saved: {path}')
